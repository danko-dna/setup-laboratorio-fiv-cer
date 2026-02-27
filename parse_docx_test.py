import sys
import docx

def analyze_docx(file_path):
    print(f"Analyzing {file_path}...")
    try:
        doc = docx.Document(file_path)
    except Exception as e:
        print(f"Error reading docx: {e}")
        return

    print(f"Total paragraphs: {len(doc.paragraphs)}")
    for i, p in enumerate(doc.paragraphs[:10]):
        print(f"P{i}: {p.text.strip()}")
    
    print(f"\nTotal tables: {len(doc.tables)}")
    for t_idx, table in enumerate(doc.tables):
        print(f"--- Table {t_idx+1} ({len(table.rows)} rows, {len(table.columns)} cols) ---")
        for r_idx, row in enumerate(table.rows):
            row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            print(f"Row {r_idx}: {row_data}")
        print("="*40)

if __name__ == "__main__":
    if len(sys.argv) > 1:
        analyze_docx(sys.argv[1])
    else:
        print("Provide a docx file path")
