import docx
doc = docx.Document("Tabla 23-12-25.docx")
print("Paragraphs:")
for p in doc.paragraphs:
    if p.text.strip(): print(f"- '{p.text.strip()}'")

print("\nTable 0, Row 0 cells:")
if len(doc.tables) > 0:
    for cell in doc.tables[0].rows[0].cells:
        if cell.text.strip(): print(f"- '{cell.text.strip()}'")
