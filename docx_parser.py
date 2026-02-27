import docx
import pandas as pd
import io

def parse_docx(file_stream, filename_fallback=""):
    """
    Lee un archivo docx aportado por el usuario y extrae sus tablas a DataFrames.
    El file_stream suele ser el resultado de st.file_uploader.
    Retorna (fecha_str, df_punciones, df_uso_interno, df_transferencias)
    """
    doc = docx.Document(file_stream)
    
    # Helper para identificar si un string parece una fecha
    def is_date_string(s):
        s_upper = s.upper()
        meses = ['ENERO', 'FEBRERO', 'MARZO', 'ABRIL', 'MAYO', 'JUNIO', 'JULIO', 'AGOSTO', 'SEPTIEMBRE', 'OCTUBRE', 'NOVIEMBRE', 'DICIEMBRE']
        dias = ['LUNES', 'MARTES', 'MIERCOLES', 'MIÉRCOLES', 'JUEVES', 'VIERNES', 'SABADO', 'SÁBADO', 'DOMINGO']
        has_month = any(m in s_upper for m in meses)
        has_day = any(d in s_upper for d in dias)
        has_year = '202' in s_upper
        return (has_month and has_year) or (has_day and has_month)
        
    fecha_str = "Fecha No Encontrada"
    candidates = []

    # 1. Buscar en Headers del documento (Si lo pusieron como encabezado)
    for section in doc.sections:
        for p in section.header.paragraphs:
            if len(p.text.strip()) > 5:
                candidates.append(p.text.strip())
                
    # 2. Buscar en Párrafos normales
    for p in doc.paragraphs:
        if len(p.text.strip()) > 5:
            candidates.append(p.text.strip())
            
    # 3. Buscar en la primera celda de la primera tabla (Por si lo pusieron dentro de la tabla)
    if len(doc.tables) > 0:
        for row in doc.tables[0].rows[:2]:
            for cell in row.cells:
                if len(cell.text.strip()) > 5:
                    candidates.append(cell.text.strip())
                    
    # Evaluar candidatos y elegir el primero que parezca fecha
    for cand in candidates:
        if is_date_string(cand):
            fecha_str = cand
            break
            
    # Fallback clásico (El primer candidato) si no encontramos una fecha clara
    if fecha_str == "Fecha No Encontrada" and candidates:
        fecha_str = candidates[0]

    # Fallback final al nombre de archivo si todo el doc está vacío de texto suelto
    if fecha_str == "Fecha No Encontrada" and filename_fallback:
        import re
        fecha_str = re.sub(r'\.docx$|\.pdf$', '', filename_fallback, flags=re.IGNORECASE)
        
    df_punciones = pd.DataFrame()
    df_uso_interno = pd.DataFrame()
    df_transferencias = pd.DataFrame()
    
    for i, table in enumerate(doc.tables):
        data = []
        for row in table.rows:
            # En vez de matar los saltos de línea con espacios, limpiamos duplicados
            # pero mantenemos los saltos simples para que el MultiCell los dibuje apilados
            row_data = []
            for cell in row.cells:
                raw_text = cell.text.strip()
                # Limpiar múltiples \n seguidos
                import re
                clean_text = re.sub(r'\n+', '\n', raw_text)
                row_data.append(clean_text)
            data.append(row_data)
            
        if not data:
            continue
            
        # Asumimos que la primera fila es el header
        headers = data[0]
        rows = data[1:]
        
        headers_lower = [str(h).lower() for h in headers]
        
        # Validar de qué tabla se trata observando la primera celda y columnas clave
        if 'hora' in headers_lower[0] and any('n° fol' in h for h in headers_lower):
            df_punciones = pd.DataFrame(rows, columns=headers)
        elif 'hora' in headers_lower[0] and any('desvitri' in h or 'ovos' in h for h in headers_lower):
            df_uso_interno = pd.DataFrame(rows, columns=headers)
        elif len(headers) >= 5 and 'hora' in headers_lower[0] and not any('semen' in h for h in headers_lower) and not any('magenta' in h for h in headers_lower):
            # Transferencias suele ser la más corta o decir DETALLE, jamás tiene Semen o Magenta
            df_transferencias = pd.DataFrame(rows, columns=headers)
            
    # Post-Procesamiento: Remover DECU OVO de la tabla de Uso Interno Lab
    if not df_uso_interno.empty:
        cols_drop = [c for c in df_uso_interno.columns if 'decu' in str(c).lower() and 'ovo' in str(c).lower()]
        if cols_drop:
            df_uso_interno = df_uso_interno.drop(columns=cols_drop)
            
    return fecha_str, df_punciones, df_uso_interno, df_transferencias
